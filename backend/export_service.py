"""
Export Service for CheckMate LMS.
Handles the generation of professionally formatted Word documents.
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

class ExportService:
    @staticmethod
    def generate_assessment_docx(title, questions):
        """
        Generates a Word document with MCQ and TF parts.
        """
        doc = Document()

        # 1. Page Header
        header = doc.add_heading(title, 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 2. Extract and Filter Parts
        # Handles both 'part' key and 'questionType' inference
        mcqs = [q for q in questions if q.get('part') == 1 or q.get('questionType') == 'MCQ']
        tfs = [q for q in questions if q.get('part') == 2 or q.get('questionType') == 'TF']

        # 3. Part 1: Multiple Choice
        if mcqs:
            doc.add_heading('PART 1: MULTIPLE CHOICE', level=1)
            for i, q in enumerate(mcqs, 1):
                # Add Question Text
                p = doc.add_paragraph(style='List Number')
                text = q.get('questionText') or q.get('text') or "Missing Question"
                run = p.add_run(text)
                run.bold = True
                
                # Add Options A, B, C, D
                options = q.get('options') or []
                for idx, opt in enumerate(options):
                    letter = chr(65 + idx) # A, B, C, D...
                    doc.add_paragraph(f"{letter}) {opt}", style='List Bullet')

        # 4. Part 2: True or False
        if tfs:
            doc.add_heading('PART 2: TRUE OR FALSE', level=1)
            # Continue numbering from the end of MCQ list
            start_num = len(mcqs) + 1
            for i, q in enumerate(tfs, start_num):
                # Add Question Text
                p = doc.add_paragraph(style='List Number')
                text = q.get('questionText') or q.get('text') or "Missing Question"
                run = p.add_run(text)
                run.bold = True
                
                # Standard T/F options
                doc.add_paragraph("A) True", style='List Bullet')
                doc.add_paragraph("B) False", style='List Bullet')

        # 5. Save to Memory Buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
